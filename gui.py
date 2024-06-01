from tkinter import *
import tkinter
from tkinter import filedialog
from lxml import etree
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter import messagebox
from PIL import Image, ImageTk
from torchvision.transforms import ToTensor
from bttr.lit_bttr import LitBTTR
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
import docx2pdf
import pytesseract



def imresize(im, sz):
    pil_im = Image.fromarray(im)
    return np.array(pil_im.resize(sz))


def resize(w_box, h_box, pil_image):
    w, h = pil_image.size
    f1 = 1.0 * w_box / w
    f2 = 1.0 * h_box / h
    factor = min([f1, f2])
    width = int(w * factor)
    height = int(h * factor)
    return pil_image.resize((width, height), Image.ANTIALIAS)



def choosepic():
    global Flag
    path_ = askopenfilename()
    path.set(path_)
    global img_open
    img_open = Image.open(e1.get()).convert('L')
    new_width = 400
    new_height = 200
    img_open_resized = img_open.resize((new_width, new_height), Image.ANTIALIAS)
    img = ImageTk.PhotoImage(img_open_resized)
    l1.config(image=img)
    l1.place(relx=0.5, y=360, anchor=CENTER)
    l1.image = img  # keep a reference
    Flag = False




def trans1():
    global prediction_string,img
    if Flag:
        messagebox.showerror(title='Error', message='No Image')
        return None
    else:
        # 替换模型加载代码
        print("Image path:", e1.get())
        ckpt = 'E:/BTTR2/pretrained-2014.ckpt'
        model = LitBTTR.load_from_checkpoint(ckpt)

        # 转换图像为模型可接受的格式
        img_pil = Image.open(e1.get()).convert('L')  # 转换为灰度图像
        img_bw = img_pil.point(lambda x: 255 if x > 128 else 0, '1')  # 二值化，将像素值大于128的设为255，小于128的设为0
        img_bw_inverted = Image.eval(img_bw, lambda x: 255 - x)  # 反转图像，将黑色变为白色，白色变为黑色
        img_bw_resized = img_bw_inverted.resize((400, 200), Image.ANTIALIAS)  # 调整大小
        img_tensor = ToTensor()(img_bw_resized)  # 转换为张量

        print("Image converted successfully!")
        print("Image tensor shape:", img_tensor.shape)
        # 进行 beam search
        hyp = model.beam_search(img_tensor)
        hyp=hyp.replace(" ", "")
        # 获取识别的 LaTeX 公式
        prediction_string = hyp
        e2 = Text(font=('Arial', 25), wrap='word', height=4, width=50)
        e2.place(relx=0.2, y=550)

        # 在原始窗口中显示 LaTeX 公式
        e2.delete(1.0, 'end')  # 清除原始窗口中的文本
        e2.insert('1.0', prediction_string)  # 在原始窗口中显示 LaTeX 公式
        return prediction_string


def trans2():
    global prediction_string
    if Flag:
        print(messagebox.showerror(title='Error', message='No Image'))
    else:
        fig = plt.figure(figsize=(3, 3))
        ax = fig.add_subplot(111)
        ax.set_xticks([])
        ax.set_yticks([])
        ax.spines['left'].set_color('none')
        ax.spines['right'].set_color('none')
        ax.spines['bottom'].set_color('none')
        ax.spines['top'].set_color('none')
        prediction_image = '$' + prediction_string + '$'
        ax.text(0.1, 0.95, prediction_image, fontsize=28)
        plt.show()
def saveClick():
    path = filedialog.asksaveasfilename(filetypes=[('Word 文档','*.docx')])
    if not path:  # 用户取消了保存操作
        return
    # path = asksaveasfilename(filetypes=[('*.txt', '.txt')])
    # with open(path, 'w+') as fb:
    #     fb.write(prediction_string)
    doc = Document()
    paragraph = doc.add_paragraph()

    # 插入公式
    run = paragraph.add_run()
    run.font.size = Pt(12)
    run.font.name = 'Cambria Math'
    # run._element.rPr.rFonts.set(nsdecls('w'), 'Cambria Math')

    mathml_element = etree.Element('{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
    mathml_element.append(etree.Element('{http://schemas.openxmlformats.org/officeDocument/2006/math}r'))
    mathml_element[0].append(etree.Element('{http://schemas.openxmlformats.org/officeDocument/2006/math}t'))
    mathml_element[0][0].text =prediction_string
    run._r.append(etree.fromstring(etree.tostring(mathml_element)))
    doc.add_paragraph("")
    doc.save(path)
#创建窗口实例
root = Tk()
#窗口大小
root.geometry('800x600')
#窗口标题
root.title('LaTeX Formula')

Flag = True
#菜单栏
menubar = Menu(root)
filemenu = Menu(menubar, tearoff=0)
menubar.add_cascade(label='Start', font=("Times New Roman", 15, 'bold'), menu=filemenu)
filemenu.add_command(label='Load', font=("Times New Roman", 15, 'bold'), command=choosepic)
#filemenu.add_command(label='Save', font=("Times New Roman", 15, 'bold'), command=saveClick)
filemenu.add_separator()
filemenu.add_command(label='Exit', font=("Times New Roman", 15, 'bold'), command=root.quit)

path = StringVar()
bu1 = Button(root, text='Start Recognization', font=('Times New Roman',18 , 'bold'), width=20, height=3, command=trans1)
bu2 = Button(root, text='Show Formula', font=('Times New Roman', 18, 'bold'), width=18, height=3, command=trans2)
bu3 = Button(root, text='Save Formula', font=('Times New Roman', 18, 'bold'), width=18, height=3, command=saveClick)
e1 = Entry(root, state='readonly', text=path)
title = Label(root, text='HMER Tool V1.0', font=('Times New Roman', 40, 'bold'), width=25, height=2)
title_1 = Label(root, text='Your image:', font=('Times New Roman', 25))
title_2 = Label(root, text='Result:', font=('Times New Roman', 25))

title.place(relx=0.5, y=40, anchor=CENTER)
bu1.place(relx=0.2, y=140, anchor=CENTER)
bu2.place(relx=0.5, y=140, anchor=CENTER)
bu3.place(relx=0.8,y=140,anchor=CENTER)
title_2.place(relx=0.1, y=500)
title_1.place(relx=0.1, y=200)

l1 = Label(root)
l1.place(relx=0.05, y=250)
l2 = Label(root)
l2.place(relx=0.55, y=250)
img_trans_show = Label(root)
img_trans_show.place(x=550, y=150)
root.config(menu=menubar)

root.mainloop()